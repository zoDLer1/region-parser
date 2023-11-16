from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dataclasses import dataclass
from configuration import REGION_IMAGES_SELL_CONFIG, REGION_IMAGES_SELLS_CONFIG, REGION_IMAGES_ROW_CONFIG, REGION_INFO_SELLS_CONFIG, SECTION_CONFIG


@dataclass
class Region:
    name: str
    status: str
    center: str
    square: int
    population: int
    region_url: str


class Docx_Regions:

    def set_section_config(self, section):
        margin = SECTION_CONFIG.get('margin', None)
        if not margin:
            return
        section.left_margin = margin.get('left', Mm(0))
        section.right_margin = margin.get('right', Mm(0))
        section.top_margin = margin.get('top', Mm(0))
        section.bottom_margin = margin.get('bottom', Mm(0))

    def __init__(self) -> None:
        self.document = Document()
        self.set_section_config(self.document.sections[0])

    region_config = REGION_IMAGES_SELL_CONFIG

    region_images_row_config = REGION_IMAGES_ROW_CONFIG

    region_images_cells_config = REGION_IMAGES_SELLS_CONFIG

    region_info_cells_config = REGION_INFO_SELLS_CONFIG

    def get_style(self):
        return self.region_config['style']

    def add_region_images(self, region: Region):
        table = self.document.add_table(rows=1, cols=2, style='Table Grid')
        for row_index, row in enumerate(table.rows):
            row.height = self.region_images_row_config[row_index]['height']
            for cell_index, cell in enumerate(row.cells):
                cell._tc.get_or_add_tcPr().append(
                    self.get_style()
                )
                paragraph = cell.add_paragraph()
                paragraph.alignment = self.region_images_cells_config[cell_index]['alignment']
                run = paragraph.add_run()
                run.add_picture(self.region_images_cells_config[cell_index]['image_path'], width=self.region_images_cells_config[
                                cell_index]['image_width'], height=self.region_images_cells_config[cell_index]['image_height'])
        end = self.document.add_paragraph(f'{region.status} {region.name}')
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_region_info(self, region: Region):
        table = self.document.add_table(rows=4, cols=2, style='Table Grid')
        table.style = 'Light Shading Accent 1'
        first_column, second_column = table.rows[0].cells[:2]
        first_column.merge(second_column)  # merge columns
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                paragraph = cell.add_paragraph(self.region_info_cells_config[row_index][cell_index]['info'](region))
                paragraph.alignment = self.region_info_cells_config[row_index][cell_index]['alignment']
        self.document.add_paragraph()

    def save(self, path):
        self.document.save(path)



